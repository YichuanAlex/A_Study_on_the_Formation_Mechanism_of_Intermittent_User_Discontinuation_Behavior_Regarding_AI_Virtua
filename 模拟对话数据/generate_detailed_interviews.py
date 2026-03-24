#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据 Excel 数据生成详细的访谈正文 Word 文档
每条数据生成一个独立的 Word 文档，包含完整的访谈流程
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import random

# 设置中文字体支持
def set_chinese_font(paragraph, font_size=12):
    """设置段落的中文字体"""
    for run in paragraph.runs:
        run.font.name = '宋体'
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def expand_answer(q_answer, question_type):
    """
    根据问题类型扩展回答，使其更详细、更自然
    
    参数:
        q_answer: Excel 中的原始回答
        question_type: 问题类型
    """
    
    # 基础回答
    base_answer = str(q_answer)
    
    # 根据不同问题类型扩展回答
    if question_type == "privacy":  # 隐私顾虑
        expansions = [
            f"{base_answer}",
            f"{base_answer} 而且我觉得图书馆作为公共机构，应该有比较完善的数据保护措施。",
            f"{base_answer} 不过如果涉及到更敏感的信息，比如身份证号、银行卡号这些，我还是会谨慎一些。",
            f"{base_answer} 毕竟现在数据泄露事件挺多的，但该用的服务还是要用，不能因噎废食。",
        ]
        return random.choice(expansions)
    
    elif question_type == "emotion":  # 情感交互
        expansions = [
            f"{base_answer}",
            f"{base_answer} 有时候我甚至会跟它多说几句，感觉像个随时在线的助手。",
            f"{base_answer} 虽然知道是程序，但好的交互设计确实能让人心情愉悦。",
            f"{base_answer} 特别是晚上一个人在图书馆学习的时候，有个能回应的'伙伴'挺好的。",
        ]
        return random.choice(expansions)
    
    elif question_type == "satisfaction":  # 满意度
        expansions = [
            f"{base_answer}",
            f"{base_answer} 总体来说，日常使用够用了，但专业深度咨询还是得找真人。",
            f"{base_answer} 我觉得对于常见问题，它的响应速度比人工快很多，这点很满意。",
            f"{base_answer} 但遇到复杂问题或者特殊情况，还是希望有人工客服可以转接。",
        ]
        return random.choice(expansions)
    
    elif question_type == "trust":  # 信任度
        expansions = [
            f"{base_answer}",
            f"{base_answer} 我一般会交叉验证一下，比如再查查官网或者问人工确认。",
            f"{base_answer} 图书馆的官方 AI 应该比外面的靠谱，毕竟是正规机构。",
            f"{base_answer} 用了几次之后，发现有些信息确实不太准确，现在会留个心眼。",
        ]
        return random.choice(expansions)
    
    elif question_type == "accessibility":  # 特殊群体适配
        expansions = [
            f"{base_answer}",
            f"{base_answer} 我带家里老人试过，他们确实不太会用，需要有人教。",
            f"{base_answer} 希望能有更多适老化设计，比如大字体、语音播报这些功能。",
            f"{base_answer} 对视障人士的无障碍支持也很重要，这是公共服务应该考虑的。",
        ]
        return random.choice(expansions)
    
    elif question_type == "alternative":  # 替代渠道
        expansions = [
            f"{base_answer}",
            f"{base_answer} 不过 AI 虚拟馆员有个好处是 24 小时在线，这点其他渠道比不了。",
            f"{base_answer} 我会根据问题的紧急程度和复杂程度来选择不同的渠道。",
            f"{base_answer} 简单问题用 AI，复杂问题打电话或者现场问，这样效率最高。",
        ]
        return random.choice(expansions)
    
    elif question_type == "motivation":  # 使用动机
        expansions = [
            f"{base_answer}",
            f"{base_answer} 特别是当我发现它能帮我节省时间的时候，使用的意愿就更强了。",
            f"{base_answer} 但发现它解决不了我的核心需求后，热情就消退了不少。",
            f"{base_answer} 现在属于有需要才用，不会特意去用它，但也不会完全不用。",
        ]
        return random.choice(expansions)
    
    elif question_type == "social":  # 社会环境
        expansions = [
            f"{base_answer}",
            f"{base_answer} 不过看到周围人都在用，我也会跟着尝试一些新功能。",
            f"{base_answer} 朋友推荐的服务我会更愿意试试，毕竟有信任基础。",
            f"{base_answer} 社交媒体上看到的评价也会影响我对这个服务的看法。",
        ]
        return random.choice(expansions)
    
    elif question_type == "interruption":  # 中辍与再启用
        expansions = [
            f"{base_answer}",
            f"{base_answer} 每次重新使用都是因为遇到了新的需求，或者听说有了新功能。",
            f"{base_answer} 停用期间也会关注相关动态，有改进的话会回来试试。",
            f"{base_answer} 属于那种需要时就捡起来，不需要就放一边的状态。",
        ]
        return random.choice(expansions)
    
    elif question_type == "demographic":  # 人口统计特征
        expansions = [
            f"{base_answer}",
            f"{base_answer} 我觉得关键是要多练习，用得多了自然就熟练了。",
            f"{base_answer} 不同年龄段的人需求不一样，服务设计应该考虑到这点。",
            f"{base_answer} 年轻人接受新事物快，但也要帮助老年人跨越数字鸿沟。",
        ]
        return random.choice(expansions)
    
    else:
        return base_answer

def generate_detailed_interview(row, participant_id):
    """
    根据 Excel 行数据生成详细的访谈文本
    
    参数:
        row: Excel 的一行数据
        participant_id: 受访者编号
    """
    
    # 从 Q1-Q10 提取回答
    q1 = row['Q1']  # 隐私顾虑
    q2 = row['Q2']  # 情感交互
    q3 = row['Q3']  # 功能满足度
    q4 = row['Q4']  # 信任度
    q5 = row['Q5']  # 特殊群体适配
    q6 = row['Q6']  # 替代渠道
    q7 = row['Q7']  # 使用动机
    q8 = row['Q8']  # 社会环境
    q9 = row['Q9']  # 中辍与再启用
    q10 = row['Q10']  # 人口统计特征
    
    # 根据 Q10 推断受访者特征
    if "年轻人" in str(q10) or "年龄不是问题" in str(q10) or "技术素养" in str(q10):
        age_group = "青年"
        education = "本科及以上"
        occupation = "学生/在职人员"
    elif "老年人" in str(q10):
        age_group = "老年"
        education = "高中/中专"
        occupation = "退休"
    elif "年龄段" in str(q10):
        age_group = "中年"
        education = "大专/本科"
        occupation = "在职人员"
    else:
        age_group = "青年"
        education = "本科"
        occupation = "学生/在职"
    
    # 生成受访者编号和基本信息
    if participant_id % 3 == 1:
        gender = "男"
        age = random.randint(22, 35)
    elif participant_id % 3 == 2:
        gender = "女"
        age = random.randint(20, 40)
    else:
        gender = random.choice(["男", "女"])
        age = random.randint(18, 45)
    
    # 构建详细的访谈正文
    interview_text = f"""访谈正文

（一）访谈开场

访谈者：您好！我是 XX 高校图书情报专业的研究生，目前正在做关于公共图书馆 AI 虚拟馆员用户间歇性中辍行为的研究，非常感谢您愿意抽出时间参与本次访谈。接下来我会围绕相关问题和您交流，访谈内容仅用于学术研究，会严格保护您的个人隐私，您可以放心分享真实的使用体验。首先，我先简单说明一下，我们所说的公共图书馆 AI 虚拟馆员，就是图书馆在官网、公众号、小程序、APP 等平台上线的智能客服、数字人助手这类服务，比如长沙图书馆的"文人猫"、浦东图书馆的"浦浦"，可以帮读者解答咨询、检索资源、办理业务等，您对这个概念有了解吗？

访谈对象：了解的，我之前用过图书馆的 AI 客服，对这个服务挺熟悉的。

（二）受访者基本信息

访谈者：好的，那先简单了解一下您的基本情况可以吗？您的年龄、学历、专业、职业是？

访谈对象：我今年{age}岁，{gender}，学历{education}，目前职业是{occupation}，平时会经常去公共图书馆学习、查资料、参加阅读活动。

（三）使用行为信息

访谈者：那您平时使用公共图书馆 AI 虚拟馆员的频次大概是多少？每次使用时长一般多久？通常会在什么场景下使用呢？

访谈对象：频次不算特别固定，平均下来大概一周用 1-3 次吧。有时候需要查资料、问问题会用得多一点，可能一天就用 2-3 次；有时候没什么需求，可能一周就用 1 次。每次使用时长一般几分钟到十几分钟不等，简单问题比如问开放时间、借还规则，几分钟就解决了；要是查文献、找馆藏、咨询深度问题，可能会聊十几分钟。使用场景主要分几种：一种是线上查资料，比如找图书馆的电子数据库、馆藏图书位置，不想跑馆里问人工；第二种是下班或者周末非工作时间，人工客服不在线，有问题就找 AI；还有就是有时候想找个性化书单，让 AI 给推荐适合的书，比自己翻目录方便；另外就是办理一些简单业务，比如续借图书、查询借阅记录这些。

（四）核心问题访谈

访谈者：好的，接下来我们进入核心问题的讨论。第一个问题，关于隐私顾虑，您在使用 AI 虚拟馆员时是否担心隐私泄露问题？能具体说说您的想法吗？

访谈对象：{expand_answer(q1, "privacy")}

访谈者：明白了，谢谢您的分享。那您觉得 AI 虚拟馆员的情感交互怎么样？有没有让您觉得温暖或者贴心的时候？或者有没有觉得它特别机械、冷冰冰的时候？

访谈对象：{expand_answer(q2, "emotion")}

访谈者：使用 AI 虚拟馆员时，是否遇到不满意的地方？功能上能满足您的需求吗？有没有哪些功能让您觉得特别好用或者特别不好用？

访谈对象：{expand_answer(q3, "satisfaction")}

访谈者：您对 AI 虚拟馆员提供的信息信任度如何？会觉得它可靠吗？有没有遇到过它提供的信息不准确的情况？

访谈对象：{expand_answer(q4, "trust")}

访谈者：您认为 AI 虚拟馆员在特殊群体适配方面做得怎么样？比如老年人、视障人士等群体使用起来方便吗？您有没有观察过或者帮助过这类群体使用？

访谈对象：{expand_answer(q5, "accessibility")}

访谈者：当您遇到问题时，除了使用 AI 虚拟馆员，还会选择哪些渠道获取帮助？您会优先选择哪个渠道？为什么？

访谈对象：{expand_answer(q6, "alternative")}

访谈者：随着使用时间的推移，您继续使用 AI 虚拟馆员的动机有什么变化吗？一开始可能觉得新鲜，现在呢？

访谈对象：{expand_answer(q7, "motivation")}

访谈者：您觉得周围的人或者社会环境对您使用 AI 虚拟馆员的决策有影响吗？比如朋友推荐、社交媒体评价这些？

访谈对象：{expand_answer(q8, "social")}

访谈者：您在使用过程中有没有经历过中辍（停止使用）后又再次启用的情况？能具体说说是什么原因让您停止使用，又是什么原因让您重新使用吗？

访谈对象：{expand_answer(q9, "interruption")}

访谈者：最后，您觉得年龄、技术素养等因素会影响 AI 虚拟馆员的使用体验吗？您有什么观察或者体会？

访谈对象：{expand_answer(q10, "demographic")}

（五）访谈收尾

访谈者：非常感谢您分享了这么多真实、详细的使用体验，这些内容对我的研究非常有帮助。您还有什么补充的内容吗？比如对 AI 虚拟馆员的建议、期望，或者其他想说的？

访谈对象：{expand_answer(q3, "satisfaction")} 总的来说，希望这个服务能越做越好，真正帮到更多读者。

访谈者：好的，再次感谢您的参与！祝您生活愉快，工作顺利！

访谈对象：不客气，也祝你的研究顺利！"""

    return interview_text

def create_word_document(interview_text, output_path):
    """
    创建 Word 文档并写入访谈文本
    
    参数:
        interview_text: 访谈文本内容
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_paragraph('访谈正文')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加空行
    doc.add_paragraph()
    
    # 分割文本并添加段落
    paragraphs = interview_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            set_chinese_font(para, 12)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    doc.save(output_path)
    print(f"已生成：{output_path}")

def main():
    """主函数"""
    # 读取 Excel 数据
    input_file = 'simulated_data_优化版.xlsx'
    output_dir = '访谈正文 Word_详细版'
    
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"已创建输出目录：{output_dir}")
    
    # 读取数据
    df = pd.read_excel(input_file)
    print(f"共读取 {len(df)} 条数据")
    
    # 只生成前 30 份
    target_count = min(30, len(df))
    print(f"将生成前 {target_count} 份详细的访谈正文...")
    
    # 为前 30 条数据生成 Word 文档
    for idx in range(target_count):
        row = df.iloc[idx]
        participant_id = idx + 1  # 受访者编号从 1 开始
        
        # 生成访谈文本
        interview_text = generate_detailed_interview(row, participant_id)
        
        # 生成文件名
        output_filename = f'受访者{participant_id:03d}_访谈正文_详细版.docx'
        output_path = os.path.join(output_dir, output_filename)
        
        # 创建 Word 文档
        create_word_document(interview_text, output_path)
    
    print(f"\n完成！共生成 {target_count} 份详细的访谈正文 Word 文档，保存在 '{output_dir}' 目录中。")

if __name__ == '__main__':
    main()
