#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据 Excel 数据生成访谈正文 Word 文档
每条数据生成一个独立的 Word 文档
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import re

# 设置中文字体支持
def set_chinese_font(paragraph, font_size=12):
    """设置段落的中文字体"""
    for run in paragraph.runs:
        run.font.name = '宋体'
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def generate_interview_text(row, participant_id):
    """
    根据 Excel 行数据生成完整的访谈文本
    
    参数:
        row: Excel 的一行数据
        participant_id: 受访者编号
    """
    
    # 从 Q1-Q10 提取回答
    q1 = str(row['Q1'])  # 隐私顾虑
    q2 = str(row['Q2'])  # 情感交互
    q3 = str(row['Q3'])  # 功能满足度
    q4 = str(row['Q4'])  # 信任度
    q5 = str(row['Q5'])  # 特殊群体适配
    q6 = str(row['Q6'])  # 替代渠道
    q7 = str(row['Q7'])  # 使用动机
    q8 = str(row['Q8'])  # 社会环境
    q9 = str(row['Q9'])  # 中辍与再启用
    q10 = str(row['Q10'])  # 人口统计特征
    
    # 生成受访者基本信息（根据 Q10 推断）
    if "年轻人" in q10 or "年龄不是问题" in q10:
        age_group = "青年"
        education = "本科及以上"
    elif "老年人" in q10:
        age_group = "老年"
        education = "高中/中专"
    else:
        age_group = "中年"
        education = "大专/本科"
    
    # 构建访谈正文
    interview_text = f"""访谈正文相关：
访谈者：您好！我是 XX 高校图书情报专业的研究生，目前正在做关于公共图书馆 AI 虚拟馆员用户间歇性中辍行为的研究，非常感谢您愿意抽出时间参与本次访谈。接下来我会围绕相关问题和您交流，访谈内容仅用于学术研究，会严格保护您的个人隐私，您可以放心分享真实的使用体验。首先，我先简单说明一下，我们所说的公共图书馆 AI 虚拟馆员，就是图书馆在官网、公众号、小程序、APP 等平台上线的智能客服、数字人助手这类服务，可以帮读者解答咨询、检索资源、办理业务等，您对这个概念有了解吗？

访谈对象：了解的，我之前用过图书馆的 AI 客服，对这个服务挺熟悉的。

（二）受访者基本信息

访谈者：好的，那先简单了解一下您的基本情况可以吗？您的年龄、学历、专业、职业是？

访谈对象：我是{age_group}人，学历{education}，平时会经常去公共图书馆学习、查资料。

（三）使用行为信息

访谈者：那您平时使用公共图书馆 AI 虚拟馆员的频次大概是多少？每次使用时长一般多久？通常会在什么场景下使用呢？

访谈对象：频次不算特别固定，平均下来大概一周用 1-3 次吧。每次使用时长一般几分钟到十几分钟不等，简单问题几分钟就解决了；要是查资料、找馆藏，可能会聊十几分钟。使用场景主要分几种：一种是线上查资料，比如找图书馆的电子数据库、馆藏图书位置；第二种是非工作时间，人工客服不在线，有问题就找 AI；还有就是想找个性化书单，让 AI 给推荐适合的书。

（四）核心问题访谈

访谈者：第一个核心问题，关于隐私顾虑，您在使用 AI 虚拟馆员时是否担心隐私泄露问题？

访谈对象：{q1}

访谈者：那您觉得 AI 虚拟馆员的情感交互怎么样？有没有让您觉得温暖或者贴心的时候？

访谈对象：{q2}

访谈者：使用 AI 虚拟馆员时，是否遇到不满意的地方？功能上能满足您的需求吗？

访谈对象：{q3}

访谈者：您对 AI 虚拟馆员提供的信息信任度如何？会觉得它可靠吗？

访谈对象：{q4}

访谈者：您认为 AI 虚拟馆员在特殊群体适配方面做得怎么样？比如老年人、视障人士等群体使用起来方便吗？

访谈对象：{q5}

访谈者：当您遇到问题时，除了使用 AI 虚拟馆员，还会选择哪些渠道获取帮助？

访谈对象：{q6}

访谈者：随着使用时间的推移，您继续使用 AI 虚拟馆员的动机有什么变化吗？

访谈对象：{q7}

访谈者：您觉得周围的人或者社会环境对您使用 AI 虚拟馆员的决策有影响吗？

访谈对象：{q8}

访谈者：您在使用过程中有没有经历过中辍（停止使用）后又再次启用的情况？能具体说说吗？

访谈对象：{q9}

访谈者：最后，您觉得年龄、技术素养等因素会影响 AI 虚拟馆员的使用体验吗？

访谈对象：{q10}

（五）访谈收尾

访谈者：非常感谢您分享了这么多真实、详细的使用体验，这些内容对我的研究非常有帮助。您还有什么补充的内容吗？

访谈对象：没有了，我觉得能说的都说到了，希望我的分享能帮到你的研究。

访谈者：好的，再次感谢您的参与！祝您生活愉快！"""

    return interview_text

def create_word_document(interview_text, output_path):
    """
    创建 Word 文档并写入访谈文本
    
    参数:
        interview_text: 访谈文本内容
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_paragraph('访谈正文')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加空行
    doc.add_paragraph()
    
    # 分割文本并添加段落
    paragraphs = interview_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            set_chinese_font(para, 12)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    doc.save(output_path)
    print(f"已生成：{output_path}")

def main():
    """主函数"""
    # 读取 Excel 数据
    input_file = 'simulated_data_优化版.xlsx'
    output_dir = '访谈正文 Word'
    
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"已创建输出目录：{output_dir}")
    
    # 读取数据
    df = pd.read_excel(input_file)
    print(f"共读取 {len(df)} 条数据")
    
    # 为每条数据生成 Word 文档
    for idx, row in df.iterrows():
        participant_id = idx + 1  # 受访者编号从 1 开始
        
        # 生成访谈文本
        interview_text = generate_interview_text(row, participant_id)
        
        # 生成文件名
        output_filename = f'受访者{participant_id:03d}_访谈正文.docx'
        output_path = os.path.join(output_dir, output_filename)
        
        # 创建 Word 文档
        create_word_document(interview_text, output_path)
        
        # 进度显示
        if (idx + 1) % 50 == 0:
            print(f"已处理 {idx + 1}/{len(df)} 条数据...")
    
    print(f"\n完成！共生成 {len(df)} 个访谈正文 Word 文档，保存在 '{output_dir}' 目录中。")

if __name__ == '__main__':
    main()
